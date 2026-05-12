import csv
from pathlib import Path


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    blocks: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if not "".join(cells).strip():
                continue
            rows.append(" | ".join(cells))
        if not rows:
            continue
        block = f"## Sheet: {ws.title}\n\n" + "\n".join(rows)
        blocks.append(block)
    return "\n\n".join(blocks)


def extract_csv(path: Path) -> str:
    with open(path, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        return "\n".join(" | ".join(row) for row in reader)


def extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")
