"""Generate six representative fixtures under ./test_data/.

Developer-only utility. Requires reportlab in addition to runtime deps:
    pip install reportlab
"""
from __future__ import annotations

import csv
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent / "test_data"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def write(path: Path, data: bytes | str) -> None:
    if isinstance(data, str):
        path.write_text(data, encoding="utf-8")
    else:
        path.write_bytes(data)
    print(f"{path.name}\t{path.stat().st_size} bytes")


def gen_intro() -> None:
    body = (
        "Welcome to Acme Health Plan.\n"
        "\n"
        "Members may visit any in-network provider without a referral.\n"
        "Emergency room visits are covered worldwide.\n"
    )
    write(OUT_DIR / "intro.txt", body)


def gen_policy() -> None:
    body = (
        "# Policy Notes\n"
        "\n"
        "- Deductible resets January 1.\n"
        "- Out-of-pocket max is $6,000 individual / $12,000 family.\n"
        "- Telehealth visits have $0 copay through Q4 2026.\n"
    )
    write(OUT_DIR / "policy.md", body)


def gen_benefits() -> None:
    rows = [
        ["service", "in_network_copay", "out_of_network_copay"],
        ["primary_care", "25", "80"],
        ["specialist", "50", "150"],
        ["urgent_care", "40", "120"],
        ["emergency_room", "300", "300"],
    ]
    path = OUT_DIR / "benefits.csv"
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        for r in rows:
            w.writerow(r)
    print(f"{path.name}\t{path.stat().st_size} bytes")


def gen_plan_summary() -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("2026 Plan Summary", level=1)
    doc.add_paragraph("This document summarizes covered benefits for the Acme PPO Gold plan.")
    doc.add_heading("Preventive Care", level=2)
    doc.add_paragraph(
        "Annual wellness exams, immunizations, and routine screenings are covered at 100% "
        "with no member cost share when received from in-network providers."
    )
    doc.add_heading("Prescription Drugs", level=2)
    table = doc.add_table(rows=5, cols=3)
    rows = [
        ["Tier", "Retail (30 day)", "Mail order (90 day)"],
        ["Generic", "$10", "$20"],
        ["Preferred brand", "$40", "$80"],
        ["Non-preferred brand", "$70", "$140"],
        ["Specialty", "$150", "N/A"],
    ]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    path = OUT_DIR / "plan_summary.docx"
    doc.save(str(path))
    print(f"{path.name}\t{path.stat().st_size} bytes")


def gen_rates() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Premiums"
    ws1.append(["tier", "monthly_premium", "employer_contribution"])
    ws1.append(["Employee only", 425, 300])
    ws1.append(["Employee + spouse", 812, 550])
    ws1.append(["Employee + family", 1180, 800])

    ws2 = wb.create_sheet("Networks")
    ws2.append(["network", "states", "hospitals_count"])
    ws2.append(["Acme National", "all", 4200])
    ws2.append(["Acme Regional", "CA,OR,WA", 380])

    path = OUT_DIR / "rates.xlsx"
    wb.save(str(path))
    print(f"{path.name}\t{path.stat().st_size} bytes")


def gen_sbc_pdf() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    path = OUT_DIR / "sbc_excerpt.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Acme PPO Gold - Benefits at a Glance", styles["Heading2"]))
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "The following table summarizes member cost-sharing for common covered services. "
            "All figures assume in-network providers unless noted.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))

    data = [
        ["Service", "Deductible Applies", "Member Cost"],
        ["Preventive care", "No", "$0"],
        ["Primary care visit", "No", "$25 copay"],
        ["Specialist visit", "No", "$50 copay"],
        ["Inpatient hospital", "Yes", "20% after deductible"],
        ["Emergency room", "Yes", "$300 copay then 20%"],
    ]
    t = Table(data, hAlign="LEFT")
    t.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ])
    )
    story.append(t)
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "Out-of-network services are subject to a separate deductible and balance billing may apply. "
            "See the full SBC for details.",
            styles["BodyText"],
        )
    )
    doc.build(story)
    print(f"{path.name}\t{path.stat().st_size} bytes")


def main() -> None:
    gen_intro()
    gen_policy()
    gen_benefits()
    gen_plan_summary()
    gen_rates()
    gen_sbc_pdf()


if __name__ == "__main__":
    main()
